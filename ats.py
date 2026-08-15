import os
import re
import io
import json
import streamlit as st
from PyPDF2 import PdfReader
from groq import Groq
from dotenv import load_dotenv
import plotly.graph_objects as go

# For DOCX/PDF exports with formatting
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle, ListStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# ------------------- Setup -------------------
load_dotenv()
API_KEY = os.getenv("GROQ_API_KEY", "")
GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")  # free-tier Groq model
if not API_KEY:
    st.warning("⚠️ GROQ_API_KEY is not set. Add it to your .env file.")
groq_client = Groq(api_key=API_KEY) if API_KEY else None

st.set_page_config(page_title="ATS Resume Reviewer", page_icon="📄", layout="wide")

# ------------------- Sidebar Branding / Help -------------------
# ------------------- Main Header -------------------
st.set_page_config(page_title="ATS Resume Reviewer", page_icon="📄", layout="wide")

st.title(" ATS Resume Reviewer")
st.caption("✨ Optimize your Resume for Applicant Tracking Systems (ATS) by comparing it against a Job Description.")

# ------------------- Instructions Section -------------------
with st.expander("ℹ️ How to Use This Tool (Click to Expand)", expanded=True):
    st.markdown(
        """
        ### 📌 Steps:
        1. **Upload your Resume (PDF format)**
        2. **Paste the Job Description** you want to match
        3. Click **Review Resume**

        ### 📊 What You'll Get:
        - ✅ ATS Match **Score**
        - 🟢 **Strengths** & ⚠️ **Weaknesses**
        - 🧩 **Missing Keywords**
        - 🔁 Suggested **Alternative Words**
        - 📝 A fully **ATS-Friendly Improved Resume** (ready for download in DOCX/PDF)

        ---
        💡 **Tip:** Use bullet points, action verbs, and clear headings in your resume to maximize ATS readability.
        """
    )


# ------------------- Prompt Template -------------------
INPUT_PROMPT = """
You are a Professional Resume Reviewer for ATS systems.
You will be provided with a Job Description and a Resume.
Your task is to analyze the Resume against the Job Description and provide feedback on how well the resume matches the job requirements.
Identify the weaknesses and strengths of the resume against the job description.
Also provide alternative words that can be used in the resume to match the Job Description.
Provide a response that is comprehensive and easy to understand for a job applicant.
Also provide a score out of 100 on how well the resume matches the job description, and how to improve the score to be competitive.
Highlight the key skills and experiences that are relevant to the job description.
Provide a well-structured and formatted resume suggestion that is ATS-friendly and matches the Job Description.
Use bullet points and proper headings. Keep each section to 200-350 words max.

IMPORTANT: You must provide the response in the following JSON format without any extra prose:

{{
   "JD Matched Score": "X%",
   "MissingKeywords": ["keyword1", "keyword2", "..."],
   "Strengths": ["strength1", "strength2", "..."],
   "Weaknesses": ["weakness1", "weakness2", "..."],
   "AlternativeWords": {{"word1": "alternative1", "word2": "alternative2", "..."}},
   "Improved Resume": "Your improved resume text here (with headings and bullet points)",
   "Profile Summary": "Your profile summary here"
}}

resume = {text}
description = {jd}
"""

# ------------------- Helpers -------------------
def input_pdf_text(file) -> str:
    """Extract text from PDF."""
    reader = PdfReader(file)
    chunks = []
    for page in reader.pages:
        txt = page.extract_text()
        if txt:
            chunks.append(txt)
    return "\n".join(chunks).strip()

def get_llm_response(prompt: str) -> str:
    """Call Groq's free-tier LLM API and return raw text."""
    if groq_client is None:
        return ""
    resp = groq_client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
    )
    return resp.choices[0].message.content or ""

def _try_json_loads(s: str):
    try:
        return json.loads(s)
    except Exception:
        return None

def clean_json_response(response_text: str):
    """Extract a JSON object from possibly noisy LLM output."""
    if not response_text:
        return None

    # 1) direct
    obj = _try_json_loads(response_text)
    if obj:
        return obj

    # 2) non-greedy braces
    matches = re.findall(r'\{.*?\}', response_text, flags=re.DOTALL)
    matches = sorted(matches, key=len, reverse=True)
    for m in matches:
        # minor cleanup
        candidate = (m.replace('\n', ' ')
                      .replace('\r', ' ')
                      .replace(', }', ' }')
                      .replace(',}', '}'))
        obj = _try_json_loads(candidate)
        if obj:
            return obj

    # 3) outermost braces
    if '{' in response_text and '}' in response_text:
        first = response_text.find('{')
        last = response_text.rfind('}')
        candidate = response_text[first:last+1]
        candidate = (candidate.replace('\n', ' ')
                            .replace('\r', ' ')
                            .replace(', }', ' }')
                            .replace(',}', '}'))
        obj = _try_json_loads(candidate)
        if obj:
            return obj

    return None

def draw_score_meter(score: int):
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=score,
        title={'text': "JD Match Score"},
        gauge={'axis': {'range': [0, 100]},
               'bar': {'color': "teal"},
               'steps': [
                   {'range': [0, 50], 'color': "lightcoral"},
                   {'range': [50, 75], 'color': "khaki"},
                   {'range': [75, 100], 'color': "lightgreen"}
               ]}
    ))
    fig.update_layout(height=300, margin=dict(t=20, b=20, l=20, r=20))
    return fig

# ---------- Formatting helpers for exports ----------
_heading_pattern = re.compile(r'^\s*(#{1,6}\s+|[A-Z][A-Za-z &/+-]{2,40}:|[A-Z ]{3,40})\s*$')
_bullet_pattern = re.compile(r'^\s*[-*•]\s+')

def parse_lines_for_formatting(text: str):
    """
    Yields tuples: ('heading'|'bullet'|'para', content)
    """
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if _heading_pattern.match(line):
            yield ('heading', line.rstrip(':').strip())
        elif _bullet_pattern.match(line):
            # remove bullet symbols
            cleaned = re.sub(r'^\s*[-*•]\s+', '', line).strip()
            yield ('bullet', cleaned)
        else:
            yield ('para', line)

# --------- DOCX generation with formatting ----------
def create_docx_formatted(resume_text: str) -> bytes:
    doc = Document()

    # Create/adjust styles
    styles = doc.styles
    if 'ATS Heading' not in [s.name for s in styles]:
        hstyle = styles.add_style('ATS Heading', WD_STYLE_TYPE.PARAGRAPH)
        hstyle.font.bold = True
        hstyle.font.size = Pt(16)
    if 'ATS Body' not in [s.name for s in styles]:
        bstyle = styles.add_style('ATS Body', WD_STYLE_TYPE.PARAGRAPH)
        bstyle.font.size = Pt(11)

    title = doc.add_paragraph("Improved Resume", style='ATS Heading')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")  # spacer

    # Build content
    list_buffer = []
    def flush_list():
        # add buffered bullets
        for item in list_buffer:
            p = doc.add_paragraph(item, style='List Bullet')
            p_format = p.paragraph_format
            p_format.space_after = Pt(2)
        list_buffer.clear()

    for kind, content in parse_lines_for_formatting(resume_text):
        if kind == 'heading':
            flush_list()
            p = doc.add_paragraph(content, style='ATS Heading')
            p_format = p.paragraph_format
            p_format.space_before = Pt(10)
            p_format.space_after = Pt(4)
        elif kind == 'bullet':
            list_buffer.append(content)
        else:  # paragraph
            flush_list()
            p = doc.add_paragraph(content, style='ATS Body')
            p_format = p.paragraph_format
            p_format.space_after = Pt(4)

    flush_list()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# --------- PDF generation with formatting ----------
def create_pdf_formatted(resume_text: str) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=54, rightMargin=54, topMargin=54, bottomMargin=54,
        title="Improved Resume"
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'ATS_Title',
        parent=styles['Heading1'],
        alignment=TA_CENTER,
        spaceAfter=12
    )
    heading_style = ParagraphStyle(
        'ATS_Heading',
        parent=styles['Heading2'],
        spaceBefore=10,
        spaceAfter=4
    )
    body_style = ParagraphStyle(
        'ATS_Body',
        parent=styles['BodyText'],
        spaceAfter=4
    )

    flow = [Paragraph("Improved Resume", title_style), Spacer(1, 6)]

    bullets = []
    def flush_bullets():
        nonlocal bullets
        if bullets:
            flow.append(ListFlowable(
                [ListItem(Paragraph(b, body_style)) for b in bullets],
                bulletType='bullet', start='-'
            ))
            flow.append(Spacer(1, 6))
            bullets = []

    for kind, content in parse_lines_for_formatting(resume_text):
        if kind == 'heading':
            flush_bullets()
            flow.append(Paragraph(content, heading_style))
        elif kind == 'bullet':
            bullets.append(content)
        else:
            flush_bullets()
            flow.append(Paragraph(content, body_style))

    flush_bullets()
    doc.build(flow)
    buffer.seek(0)
    return buffer.getvalue()

# ------------------- Main UI -------------------
st.title(" ATS Resume Reviewer")
st.caption("✨ Compare your Resume against a Job Description and optimize it for ATS success.")

col1, col2 = st.columns([1, 1])
with col1:
    jd = st.text_area("📝 Job Description", height=300, placeholder="Paste the JD here...")
with col2:
    uploaded_file = st.file_uploader("📂 Upload your Resume (PDF)", type=["pdf"])

# Form to allow Enter-key submission
with st.form("review_form", clear_on_submit=False):
    submitted = st.form_submit_button(" Review Resume", use_container_width=True)

if submitted:
    if not API_KEY:
        st.error("No API key found. Please set GROQ_API_KEY in .env")
    elif not jd or not uploaded_file:
        st.error("Please upload a resume and provide a job description.")
    else:
        with st.spinner("🔎 Extracting resume text..."):
            resume_text = input_pdf_text(uploaded_file)

        if not resume_text.strip():
            st.error("Could not extract any text from the PDF. Please try another file.")
        else:
            with st.expander("Preview: Extracted Resume Text", expanded=False):
                st.write(resume_text[:3000] + ("..." if len(resume_text) > 3000 else ""))

            prompt = INPUT_PROMPT.format(text=resume_text, jd=jd)

            with st.spinner("🤖 Analyzing with Groq..."):
                raw_response = get_llm_response(prompt)

            parsed = clean_json_response(raw_response)

            if not parsed:
                st.error("❌ Could not parse a valid JSON response.")
                with st.expander("Raw Model Response"):
                    st.write(raw_response)
            else:
                # Tabs
                tabs = st.tabs(["📊 Analysis", "📝 Improved Resume", "👤 Profile Summary", "🛠 Debug"])

                with tabs[0]:
                    st.subheader("📊 Resume Analysis")

                    # Score
                    score_text = parsed.get("JD Matched Score", "0%").replace("%", "")
                    try:
                        score_val = int(re.findall(r'\d+', score_text)[0])
                    except Exception:
                        score_val = 0
                    st.plotly_chart(draw_score_meter(score_val), use_container_width=True)

                    colA, colB = st.columns(2)
                    with colA:
                        st.markdown("### ✅ Strengths")
                        strengths = parsed.get("Strengths", [])
                        st.success("\n".join([f"• {s}" for s in strengths]) or "N/A")

                        st.markdown("### 🧩 Missing Keywords")
                        missing = parsed.get("MissingKeywords", [])
                        st.warning("\n".join([f"• {k}" for k in missing]) or "N/A")

                    with colB:
                        st.markdown("### ⚠️ Weaknesses")
                        weaknesses = parsed.get("Weaknesses", [])
                        st.error("\n".join([f"• {w}" for w in weaknesses]) or "N/A")

                        st.markdown("### 🔁 Alternative Words")
                        alt = parsed.get("AlternativeWords", {})
                        if alt:
                            for w, a in alt.items():
                                st.write(f"- **{w}** → {a}")
                        else:
                            st.info("N/A")

                with tabs[1]:
                    st.subheader("📝 Improved Resume (ATS-Friendly)")
                    improved_resume = parsed.get("Improved Resume", "N/A")
                    st.markdown(improved_resume)

                    if improved_resume and improved_resume != "N/A":
                        colx, coly = st.columns(2)
                        with colx:
                            st.download_button(
                                label="⬇️ Download as Word (.docx)",
                                data=create_docx_formatted(improved_resume),
                                file_name="Improved_Resume.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                        with coly:
                            st.download_button(
                                label="⬇️ Download as PDF (.pdf)",
                                data=create_pdf_formatted(improved_resume),
                                file_name="Improved_Resume.pdf",
                                mime="application/pdf",
                                use_container_width=True
                            )

                with tabs[2]:
                    st.subheader("👤 Profile Summary")
                    st.info(parsed.get("Profile Summary", "N/A"))

                with tabs[3]:
                    st.subheader("🛠 Debug & JSON")
                    st.json(parsed)
                    with st.expander("Raw LLM Response"):
                        st.write(raw_response)
# ------------------- Export Helpers -------------------    
# (Moved above main UI for better organization)
# (Moved above main UI for better organization)